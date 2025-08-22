# src/parse_resume.py
from pathlib import Path
import re
from typing import Dict, List, Tuple
import io

# optional pdfminer fallback
try:
    from pdfminer.high_level import extract_text as pdfminer_extract
except Exception:
    pdfminer_extract = None

# optional docx (python-docx)
try:
    import docx
except Exception:
    docx = None

# PyPDF2 is the primary extractor for PDFs
import PyPDF2


def extract_text(file) -> str:
    """
    Robust extractor for UploadedFile (Streamlit) or file path string.
    - For PDF: try PyPDF2, then pdfminer if available.
    - For DOCX: use python-docx if available.
    - For TXT: decode bytes.
    Returns plain text (possibly empty string).
    """
    # Support passing path string
    if isinstance(file, (str, Path)):
        with open(str(file), "rb") as fh:
            data = fh.read()
    else:
        # Streamlit UploadedFile supports .read()
        try:
            file.seek(0)
        except Exception:
            pass
        data = file.read()

    if not data:
        return ""

    text = ""
    # work with bytes always
    b = io.BytesIO(data)

    # PDF
    if getattr(file, "name", "").lower().endswith(".pdf") or (isinstance(file, str) and str(file).lower().endswith(".pdf")):
        # PyPDF2 attempt
        try:
            b.seek(0)
            reader = PyPDF2.PdfReader(b)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            # don't crash — just continue to fallback(s)
            print(f"[extract_text] PyPDF2 error: {e}")

        # pdfminer fallback (if installed)
        if not text.strip() and pdfminer_extract is not None:
            try:
                b.seek(0)
                # pdfminer can accept a file-like object
                text = pdfminer_extract(b)
            except Exception as e:
                print(f"[extract_text] pdfminer error: {e}")

    # DOCX
    elif getattr(file, "name", "").lower().endswith(".docx") or (isinstance(file, str) and str(file).lower().endswith(".docx")):
        if docx is None:
            print("[extract_text] python-docx not installed; cannot read docx.")
            return ""
        try:
            b.seek(0)
            d = docx.Document(io.BytesIO(data))
            paras = [p.text for p in d.paragraphs if p.text]
            text = "\n".join(paras)
        except Exception as e:
            print(f"[extract_text] docx error: {e}")
            return ""

    # TXT
    elif getattr(file, "name", "").lower().endswith(".txt") or (isinstance(file, str) and str(file).lower().endswith(".txt")):
        try:
            if isinstance(data, bytes):
                text = data.decode("utf-8", errors="ignore")
            else:
                text = str(data)
        except Exception as e:
            print(f"[extract_text] txt decode error: {e}")
            return ""

    else:
        # unknown extension — try decoding as text
        try:
            text = data.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    return (text or "").strip()


def load_text(file_path: str) -> str:
    """
    Legacy helper: read from a local path (keeps compatibility with older code).
    """
    p = Path(file_path)
    ext = p.suffix.lower()
    if ext == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        if pdfminer_extract is None:
            raise RuntimeError("pdfminer.six not installed. Install or use .txt/.docx")
        return pdfminer_extract(str(p))
    if ext == ".docx":
        if docx is None:
            raise RuntimeError("python-docx not installed. Install or use .txt/.pdf")
        d = docx.Document(str(p))
        return "\n".join([para.text for para in d.paragraphs])
    raise ValueError(f"Unsupported file type: {ext}")


def normalize(text: str) -> str:
    """
    Normalize text for matching. Keep + and # (for C++ / C#).
    """
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r"[^a-z0-9#+.\- ]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_years_experience(text: str) -> int:
    matches = re.findall(r"(\d+)\s*\+?\s*year", (text or "").lower())
    years = [int(m) for m in matches] if matches else [0]
    return max(years) if years else 0


def build_alias_map(skills_db: List[Dict]) -> Dict[str, str]:
    amap = {}
    for s in skills_db:
        canonical = s["name"].lower()
        amap[canonical] = canonical
        for a in s.get("aliases", []):
            amap[a.lower()] = canonical
    return amap


def extract_skills(text: str, skills_db: List[Dict]) -> Tuple[List[str], Dict[str, List[Tuple[int, int]]]]:
    """
    Extract skills using alias map. Returns (ordered_skills, spans).
    """
    if not text:
        return [], {}
    norm = normalize(text)
    amap = build_alias_map(skills_db)
    patterns = sorted(amap.keys(), key=lambda x: len(x), reverse=True)

    found = []
    spans = {}
    for pat in patterns:
        # make regex forgiving: allow punctuation/boundaries; case-insensitive
        regex = re.compile(rf"(?<![a-z0-9]){re.escape(pat)}(?![a-z0-9])", flags=re.IGNORECASE)
        for m in regex.finditer(norm):
            can = amap[pat]
            found.append(can)
            spans.setdefault(can, []).append((m.start(), m.end()))

    # dedupe preserve order
    seen = set()
    ordered = []
    for s in found:
        if s not in seen:
            ordered.append(s)
            seen.add(s)
    return ordered, spans
